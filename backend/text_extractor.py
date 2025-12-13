import os
import PyPDF2
import docx
import chardet

# -----------------------------------------------
# Extract text from TXT
# -----------------------------------------------
def extract_text_from_txt(file_path):
    try:
        # Detect encoding first
        with open(file_path, "rb") as f:
            raw = f.read()
            result = chardet.detect(raw)
            encoding = result["encoding"] or "utf-8"

        # Try reading with detected encoding
        with open(file_path, "r", encoding=encoding, errors="ignore") as f:
            return f.read()

    except Exception as e:
        return f"ERROR: Failed to read TXT file → {e}"


# -----------------------------------------------
# Extract text from PDF
# -----------------------------------------------
def extract_text_from_pdf(file_path):
    text = ""

    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)

            # Handle password-protected PDFs
            if reader.is_encrypted:
                return "ERROR: PDF is password-protected."

            for page in reader.pages:
                extracted = page.extract_text()
                
                if extracted:
                    text += extracted + "\n\n"
                else:
                    # PDF may be scanned (image-based)
                    return "ERROR: PDF appears to be scanned (no extractable text). OCR required."

        return text.strip()

    except Exception as e:
        return f"ERROR: Failed to read PDF → {e}"


# -----------------------------------------------
# Extract text from DOCX
# -----------------------------------------------
def extract_text_from_docx(file_path):
    try:
        document = docx.Document(file_path)
        text = ""

        # Extract text from paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract from tables
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"

        return text.strip()

    except Exception as e:
        return f"ERROR: Failed to read DOCX → {e}"


# -----------------------------------------------
# Text Cleaning
# -----------------------------------------------
def clean_text(text):
    if not text or text.startswith("ERROR"):
        return text

    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove extra blank lines
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    cleaned = "\n".join(lines)

    return cleaned


# -----------------------------------------------
# Unified Extraction Function
# -----------------------------------------------
def extract_text(file_path):
    if not os.path.exists(file_path):
        return {"status": "error", "message": "File does not exist"}

    ext = file_path.split(".")[-1].lower()

    if ext == "txt":
        raw_text = extract_text_from_txt(file_path)
    elif ext == "pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext == "docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        return {"status": "error", "message": f"Unsupported file type: {ext}"}

    if isinstance(raw_text, str) and raw_text.startswith("ERROR"):
        return {"status": "error", "message": raw_text}

    cleaned = clean_text(raw_text)

    if len(cleaned.strip()) == 0:
        return {"status": "error", "message": "Extracted text is empty"}

    metadata = {
        "word_count": len(cleaned.split()),
        "char_count": len(cleaned),
    }

    return {
        "status": "success",
        "text": cleaned,
        "metadata": metadata
    }
